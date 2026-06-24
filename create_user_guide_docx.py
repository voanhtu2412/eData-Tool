from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(r"C:\Users\pc\Downloads\Huong_dan_su_dung_eData_TaxCodeCollector.docx")
SCREENSHOT = Path(r"C:\Users\pc\AppData\Local\Temp\codex-clipboard-b87effb3-3a9f-4d3c-9b65-2b915b1bdb9f.png")


def set_run_font(run, name="Calibri", size=None, color=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, size=10, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        set_cell_shading(cell, fill)


def set_table_borders(table, color="D8DEE9", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_title(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Hướng dẫn sử dụng eData")
    set_run_font(run, size=24, color="4F46E5", bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Template-based Web Scraper cho dữ liệu mã số thuế")
    set_run_font(run, size=13, color="475569")

    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(1.7)
    table.columns[1].width = Inches(4.8)
    set_table_borders(table, color="E2E8F0", size="4")
    set_cell_margins(table)
    rows = [
        ("Mục đích", "Chọn mẫu dữ liệu một lần, sau đó tự động lặp lại mẫu đó cho nhiều công ty."),
        ("Đầu ra", "DataGrid trong tool và file Excel chỉ gồm các field người dùng đã chọn."),
        ("Ngày tạo", datetime.now().strftime("%d/%m/%Y")),
    ]
    for row, (label, value) in zip(table.rows, rows):
        set_cell_text(row.cells[0], label, bold=True, fill="EEF2FF")
        set_cell_text(row.cells[1], value)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_note(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table, color="C7D2FE", size="6")
    set_cell_margins(table, top=120, bottom=120, start=160, end=160)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "EEF2FF")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    set_run_font(run, size=10.5, color="3730A3", bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(body)
    set_run_font(run, size=10.5, color="1E293B")


def add_troubleshooting_table(doc):
    doc.add_heading("7. Lỗi thường gặp và cách xử lý", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    widths = [1.75, 2.25, 2.5]
    for idx, width in enumerate(widths):
        table.columns[idx].width = Inches(width)
    set_table_borders(table)
    set_cell_margins(table)
    headers = ["Hiện tượng", "Nguyên nhân thường gặp", "Cách xử lý"]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(cell, header, bold=True, fill="E8EEF5")

    rows = [
        ("Pick Element không mở trình duyệt", "Playwright chưa cài Chromium.", "Chạy file playwright.ps1 install chromium rồi mở lại tool."),
        ("Start Crawl không có dòng dữ liệu", "URL danh sách chưa đúng hoặc website không có link công ty ở trang hiện tại.", "Kiểm tra lại URL danh sách, URL trang 2 và số trang cần lọc."),
        ("Một field bị trống", "Selector/XPath không khớp ở trang chi tiết công ty.", "Mở trang chi tiết mẫu, chọn lại field đó và bấm Pick Element lại."),
        ("Excel không có cột Tên công ty hoặc Url", "Tool chỉ xuất field đã chọn.", "Nếu cần cột đó, thêm field template tên Tên Cty hoặc Url rồi pick element tương ứng."),
        ("Không build được khi đang debug", "Visual Studio hoặc app đang khóa file .exe/.dll.", "Stop Debugging, đóng app đang chạy, sau đó build lại."),
    ]
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            set_cell_text(cell, value)


def build_doc():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_title(doc)

    doc.add_heading("1. Tool này dùng để làm gì?", level=1)
    p = doc.add_paragraph()
    p.add_run("eData là tool desktop giúp lấy dữ liệu công ty từ website mã số thuế theo cơ chế ")
    r = p.add_run("Template-based Web Scraper")
    set_run_font(r, bold=True)
    p.add_run(". Người dùng chọn mẫu dữ liệu một lần, tool sẽ tự động vào từng công ty và lấy đúng các trường đã chọn.")

    add_note(
        doc,
        "Nguyên tắc quan trọng",
        "Bảng kết quả và file Excel chỉ hiển thị các field nằm trong Template Fields. Nếu muốn có cột nào, hãy thêm field đó vào template."
    )

    doc.add_heading("2. Chuẩn bị trước khi crawl", level=1)
    add_bullets(doc, [
        "Mở tool eData - Template-based Web Scraper.",
        "Xác định một trang chi tiết công ty mẫu để pick dữ liệu, ví dụ trang có đủ Tên Cty, MST, SĐT, Người đại diện, Ngày hoạt động.",
        "Xác định URL danh sách công ty cần crawl.",
        "Nếu danh sách có phân trang, chuẩn bị URL trang 2 hoặc URL trang tiếp theo.",
    ])

    doc.add_heading("3. Cấu hình Template Fields", level=1)
    doc.add_paragraph("Template Fields là danh sách cột dữ liệu bạn muốn lấy. Mỗi dòng template gồm:")
    add_bullets(doc, [
        "FieldName: tên cột sẽ hiển thị trong DataGrid và Excel, ví dụ Tên Cty, MST, SĐT.",
        "CssSelector: selector dùng để lấy dữ liệu trên trang chi tiết.",
        "XPath: XPath dự phòng nếu CssSelector không lấy được dữ liệu.",
        "SampleValue: giá trị mẫu sau khi pick element, dùng để kiểm tra bạn đã chọn đúng phần tử.",
    ])
    doc.add_heading("Cách thêm field", level=2)
    add_numbers(doc, [
        "Bấm Add Field để thêm dòng mới.",
        "Sửa FieldName thành tên dữ liệu cần lấy.",
        "Nhập URL trang chi tiết mẫu vào ô URL dùng khi Pick Element.",
        "Chọn dòng field đó trong bảng Template Fields.",
        "Bấm Pick Element.",
        "Khi trình duyệt Playwright mở ra, click đúng phần tử chứa dữ liệu trên website.",
        "Kiểm tra SampleValue. Nếu đúng, bấm Save Template.",
    ])

    doc.add_heading("4. Crawl dữ liệu công ty", level=1)
    add_numbers(doc, [
        "Nhập URL danh sách công ty vào ô URL danh sách công ty.",
        "Nếu có phân trang, nhập URL trang 2 / trang tiếp theo.",
        "Nhập Số trang cần lọc. Ví dụ nhập 1 để lấy trang hiện tại, nhập 5 để lấy 5 trang.",
        "Bấm Start Crawl.",
        "Theo dõi Log trạng thái để biết tool đang mở công ty nào và đang lấy field nào.",
        "Nếu muốn dừng giữa chừng, bấm Stop Crawl.",
    ])

    doc.add_heading("5. Kiểm tra kết quả", level=1)
    add_bullets(doc, [
        "Mỗi công ty tương ứng một dòng trong bảng kết quả.",
        "Mỗi field đã chọn trong template tương ứng một cột.",
        "Nếu bạn chỉ chọn Tên Cty, MST, SĐT, Người đại diện thì bảng chỉ có đúng 4 cột đó.",
        "Nếu một ô bị trống, thường là selector của field đó không còn khớp với trang chi tiết.",
    ])

    doc.add_heading("6. Export Excel", level=1)
    add_numbers(doc, [
        "Sau khi crawl xong, bấm Export Excel.",
        "Chọn nơi lưu file .xlsx.",
        "Mở file Excel để kiểm tra dữ liệu.",
    ])
    add_note(
        doc,
        "Lưu ý khi export",
        "Excel chỉ xuất các field trong template, không tự thêm Url hoặc Tên công ty nếu bạn không tạo field đó."
    )

    add_troubleshooting_table(doc)

    doc.add_heading("8. Gợi ý sử dụng hiệu quả", level=1)
    add_bullets(doc, [
        "Đặt FieldName ngắn, rõ nghĩa: Tên Cty, MST, SĐT, Người đại diện, Địa chỉ.",
        "Luôn pick từ trang chi tiết công ty có đầy đủ dữ liệu nhất.",
        "Sau khi thêm field mới, nên crawl thử 1 trang trước khi crawl số lượng lớn.",
        "Nếu website đổi giao diện, cần pick lại các field bị trống.",
        "Nên lưu Excel sau mỗi phiên crawl lớn để tránh mất dữ liệu do website hoặc mạng lỗi giữa chừng.",
    ])

    if SCREENSHOT.exists():
        doc.add_section(WD_SECTION.NEW_PAGE)
        doc.add_heading("Phụ lục: Giao diện chính", level=1)
        p = doc.add_paragraph("Ảnh minh họa giao diện eData với khu vực Template Fields bên trái và cấu hình Crawl bên phải.")
        p.paragraph_format.space_after = Pt(8)
        doc.add_picture(str(SCREENSHOT), width=Inches(6.5))

    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("eData - Template-based Web Scraper")
    set_run_font(run, size=9, color="64748B")

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_doc())
